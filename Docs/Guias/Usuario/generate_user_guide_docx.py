from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "GUIA_USUARIO_ACC.md"
OUTPUT_DOCX = ROOT / "GUIA_USUARIO_ACC.docx"


IMAGE_RE = re.compile(r"^!\[(?P<alt>.*?)]\((?P<path>.*?)\)\s*$")
ORDERED_RE = re.compile(r"^(?P<num>\d+)\.\s+(?P<text>.+)$")
UNORDERED_RE = re.compile(r"^-\s+(?P<text>.+)$")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def ensure_style(document: Document, name: str, base_name: str | None = None):
    styles = document.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base_name and base_name in styles:
        style.base_style = styles[base_name]
    return style


def set_font(run, bold: bool = False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def configure_document(document: Document):
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    body = ensure_style(document, "ACCBody", "Normal")
    body.font.name = "Times New Roman"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    body.font.size = Pt(12)
    body.font.color.rgb = RGBColor(0, 0, 0)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.15

    bullet = ensure_style(document, "ACCBullet", "Normal")
    bullet.font.name = "Times New Roman"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    bullet.font.size = Pt(12)
    bullet.font.color.rgb = RGBColor(0, 0, 0)
    bullet.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bullet.paragraph_format.left_indent = Inches(0.25)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(3)

    numbered = ensure_style(document, "ACCNumbered", "Normal")
    numbered.font.name = "Times New Roman"
    numbered._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    numbered.font.size = Pt(12)
    numbered.font.color.rgb = RGBColor(0, 0, 0)
    numbered.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    numbered.paragraph_format.left_indent = Inches(0.25)
    numbered.paragraph_format.first_line_indent = Inches(-0.25)
    numbered.paragraph_format.space_after = Pt(3)

    caption = ensure_style(document, "ACCCaption", "Normal")
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(12)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(9)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def strip_code_ticks(text: str) -> str:
    return INLINE_CODE_RE.sub(lambda m: m.group(1), text)


def add_text_runs(paragraph, text: str):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    add_text_runs(paragraph, strip_code_ticks(text))


def add_body(document: Document, text: str):
    paragraph = document.add_paragraph(style="ACCBody")
    add_text_runs(paragraph, text)


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="ACCBullet")
    add_text_runs(paragraph, f"- {text}")


def add_numbered(document: Document, number: str, text: str):
    paragraph = document.add_paragraph(style="ACCNumbered")
    add_text_runs(paragraph, f"{number}. {text}")


def add_image(document: Document, image_path: Path, alt_text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.0))

    if alt_text.strip():
        caption = document.add_paragraph(style="ACCCaption")
        add_text_runs(caption, alt_text.strip())


def remove_empty_first_paragraph(document: Document):
    if len(document.paragraphs) != 1:
        return
    if document.paragraphs[0].text.strip():
        return
    p = document.paragraphs[0]._element
    p.getparent().remove(p)


def enable_update_fields_on_open(document: Document):
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_docx():
    markdown = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)
    remove_empty_first_paragraph(document)

    for raw_line in markdown:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            image_rel = unquote(image_match.group("path"))
            image_path = (ROOT / image_rel).resolve()
            if image_path.exists():
                add_image(document, image_path, image_match.group("alt"))
            else:
                add_body(document, f"[Imagen no encontrada: {image_rel}]")
            continue

        if stripped.startswith("### "):
            add_heading(document, stripped[4:], 3)
            continue

        if stripped.startswith("## "):
            add_heading(document, stripped[3:], 2)
            continue

        if stripped.startswith("# "):
            add_heading(document, stripped[2:], 1)
            continue

        ordered_match = ORDERED_RE.match(stripped)
        if ordered_match:
            add_numbered(document, ordered_match.group("num"), ordered_match.group("text"))
            continue

        unordered_match = UNORDERED_RE.match(stripped)
        if unordered_match:
            add_bullet(document, unordered_match.group("text"))
            continue

        add_body(document, stripped)

    enable_update_fields_on_open(document)
    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
